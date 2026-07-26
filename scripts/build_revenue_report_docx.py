#!/usr/bin/env python3
"""Build Word document for Dandelion revenue analysis report."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Dandelion_Revenue_Analysis_Report.docx"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Dandelion Flowers", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Revenue Analysis — Benchmarks & Alameda Targets")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    meta = doc.add_paragraph("Prepared: July 26, 2026  |  Week 4 Deliverable  |  Dandelion Flowers · Alameda, CA")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Dandelion annualizes to approximately $316,000 in revenue (2026 year-to-date pace) — "
        "squarely in the $300K–$500K healthy independent florist band. The shop is not revenue-starved; "
        "it is margin- and mix-starved."
    )

    add_table(doc,
        ["Verdict", "Metric", "Actual", "Alameda Target", "Gap"],
        [
            ["Winning", "Online AOV", "$95–$133", "$85–$115", "At or above"],
            ["Winning", "Online channel share", "49%", "42–48%", "At target"],
            ["Winning", "2026 online growth", "+41% YTD vs 2025", "Protect recovery", "Ahead"],
            ["Mixed", "Repeat rate", "21–31%", "40%+", "Below strong performer"],
            ["Losing", "POS median AOV", "$49", "$85–$115", "−$36 to −$66"],
            ["Losing", "Gross margin", "31.5%", "62–68%", "−31 pts"],
            ["Losing", "Wholesale spend / revenue", "63%", "28–35%", "~2× benchmark"],
            ["Losing", "Card attach", "0.6%", "30%+", "Near zero"],
        ],
    )

    doc.add_paragraph(
        "Market context: Alameda median household income is $137,697 (+39% vs California). "
        "Customers can afford premium floral. Online pricing proves it; the counter does not."
    )

    doc.add_heading("Two Focus Metrics — Next 8 Weeks", level=2)
    add_bullets(doc, [
        "POS median AOV: $49 → $75 by end of Q3",
        "Shop gross margin %: 31.5% → 38% by end of Q3",
    ])
    doc.add_paragraph(
        "Everything else (SEO, weddings, corporate) is secondary until these move."
    )

    doc.add_heading("1. Revenue & Shop Size", level=1)
    doc.add_heading("Annual Revenue Trend", level=2)
    add_table(doc,
        ["Year", "Bank Revenue", "SS Online", "Online %", "GP%", "Context"],
        [
            ["2023", "$333,737", "$110,123", "33%", "38.1%", "Above $300K — healthy band"],
            ["2024", "$282,530", "$107,984", "38%", "34.3%", "Declining; margin slipping"],
            ["2025", "$276,591", "$106,459", "38%", "25.8%", "Below $300K; margin crisis"],
            ["2026 YTD (7 mo)", "$184,160", "$89,835", "49%", "31.5%", "Recovering"],
            ["2026 annualized", "~$316K", "~$154K", "~49%", "~31.5%", "$300K–$500K band"],
        ],
    )

    doc.add_heading("2026 Channel Mix", level=2)
    add_table(doc,
        ["Channel", "Revenue", "Orders", "Share", "Benchmark"],
        [
            ["Square POS + invoices", "$111,578", "1,680", "60%", "Walk-in GM 40–50%"],
            ["Squarespace (total)", "$89,835", "678", "40%", "Direct web GM 45–65%"],
            ["Online share (bank)", "—", "—", "49%", "National: 42–48%"],
        ],
    )

    doc.add_heading("Seasonality", level=2)
    add_table(doc,
        ["Window", "Dandelion 2026", "Industry Note"],
        [
            ["Valentine's month (Feb)", "21% of YTD", "Peak; 20–40% above baseline"],
            ["Mother's Day month (May)", "25% of YTD", "Peak"],
            ["Feb + May combined", "~46%", "Plan inventory + retention"],
            ["Soft months (Apr, Jun)", "9–11% GP%", "Margin floor — stems not cut"],
        ],
    )

    doc.add_heading("2. AOV — Actuals vs Alameda Targets", level=1)
    add_table(doc,
        ["Segment", "Actual", "National", "Alameda Target", "Status"],
        [
            ["Square POS median", "$49", "$65–$90", "$85–$115", "Below"],
            ["Square POS mean", "$66", "$65–$90", "$85–$115", "Low end"],
            ["SS product median", "$95", "$72–$83", "$85–$115", "On target"],
            ["SS total median", "$121", "—", "$85–$115", "Above target"],
            ["Blended all channels", "$78", "$65–$90", "$85–$115", "Below"],
        ],
    )
    doc.add_paragraph(
        "Dollar impact: Alameda target blended AOV ≈ $95 vs actual $78 on ~2,358 orders YTD ≈ "
        "$40,000 left on table (annualized ~$69K). 40% of POS tickets are under $40."
    )

    doc.add_heading("Price Ladder", level=2)
    add_table(doc,
        ["Tier", "Online (working)", "POS (broken)"],
        [
            ["Entry", "Touch of Honey ~$65–68", '"20", "45 bouquet" ~$40–60'],
            ["Mid", "Love Poem, Moon ~$85–101", "Generic Bouquet"],
            ["Premium", "Darling ~$180, Great Expectations ~$309", "Rarely sold at counter"],
        ],
    )

    doc.add_heading("3. Gross Margin — Actuals vs Benchmarks", level=1)
    add_table(doc,
        ["Metric", "Dandelion", "National Healthy", "Alameda Target", "Status"],
        [
            ["Gross margin % (2026 YTD)", "31.5%", "65–72%", "62–68%", "Critical"],
            ["Gross margin % (2025)", "25.8%", "45–58%", "62–68%", "Deteriorating"],
            ["Wholesale/flowers % of revenue", "63.2%", "28–35%", "Under 40%", "~2× benchmark"],
        ],
    )

    doc.add_heading("Margin Trend", level=2)
    add_table(doc,
        ["Year", "Revenue", "GP%", "Wholesale %", "What Happened"],
        [
            ["2023", "$334K", "38.1%", "54%", "Highest revenue + best margin"],
            ["2024", "$283K", "34.3%", "59%", "Revenue −15%; margin slipping"],
            ["2025", "$277K", "25.8%", "68%", "Stem spend up; prices flat"],
            ["2026 YTD", "$184K", "31.5%", "63%", "Volume recovering; margin broken"],
        ],
    )

    doc.add_heading("Monthly Margin 2026", level=2)
    add_table(doc,
        ["Month", "Revenue", "GP%", "Read"],
        [
            ["Feb (Valentine's)", "$44,644", "42.6%", "Best month — still below 62% target"],
            ["May (Mother's Day)", "$45,339", "48.3%", "Peak revenue + best GP%"],
            ["Apr", "$19,910", "8.7%", "Revenue down; stems not cut"],
            ["Jun", "$14,582", "11.4%", "Same pattern"],
        ],
    )

    doc.add_heading("Hero Product Margins (Predicted)", level=2)
    doc.add_paragraph(
        "Estimated via weekly COGS allocation (no inventory carry-over). Full COGS, 4-week rolling rate."
    )
    add_table(doc,
        ["Product", "Revenue", "ASP", "Est. GP%"],
        [
            ["Touch of Honey", "$41,920", "$68", "30.4%"],
            ["Love Poem", "$38,287", "$91", "31.1%"],
            ["To the Moon and Back", "$36,315", "$101", "33.7%"],
            ["TLC", "$31,013", "$150", "32.5%"],
            ["XoXo", "$24,499", "$118", "33.5%"],
            ["Deep in the Woods", "$23,753", "$130", "30.4%"],
            ["The Sympathy", "$20,785", "$111", "31.5%"],
            ["Darling", "$12,723", "$190", "37.2%"],
            ["Great Expectations", "$7,719", "$309", "27.3%"],
        ],
    )
    doc.add_paragraph("All 9 heroes are below 38% gross margin on full COGS. Touch of Honey is the worst.")

    doc.add_heading("4. Customer & Retention", level=1)
    add_table(doc,
        ["Metric", "Dandelion", "Typical", "Alameda Target", "Status"],
        [
            ["Repeat rate (online email)", "20.9%", "22–35%", "40%+", "Below"],
            ["Repeat rate (Square)", "31.1%", "22–35%", "40%+", "Typical"],
            ["2026 orders from returning emails", "42.9%", "—", "—", "Improving"],
            ["Orders per customer / year", "1.56", "2.5–4.0", "3.0+", "Below"],
            ["Email capture (Square)", "7.4%", "—", "85%+ txs", "Critical"],
            ["Online payer emails", "5,986", "—", "—", "Key asset"],
        ],
    )

    doc.add_heading("Customer Value Tiers (Online)", level=2)
    add_table(doc,
        ["Tier", "Customers", "% of Base", "Note"],
        [
            ["VIP ($1,000+)", "85", "1.4%", "14.6% of SS revenue"],
            ["Regular ($300–999)", "477", "8.0%", "Sub/win-back candidates"],
            ["Occasional ($100–299)", "2,745", "45.9%", "Largest pool — occasion reminders"],
            ["Light (<$100)", "2,679", "44.8%", "One-time / holiday buyers"],
        ],
    )

    doc.add_heading("5. High-LTV Channels", level=1)
    add_table(doc,
        ["Channel", "Dandelion", "Benchmark", "Gap"],
        [
            ["Subscriptions", "4 orders YTD; $195–200/mo", "18% of online rev long-term", "Near zero activity"],
            ["Weddings", "Not tracked; invoices $165 avg", "$5,500–$8,000 contract", "No pipeline"],
            ["Sympathy", "50 orders; $6,941; $139 AOV", "31.4% of Alameda 65+", "Funeral homes not started"],
            ["Corporate", "6 company profiles", "20 accts × $185/wk ≈ $192K/yr", "Not started"],
        ],
    )

    doc.add_heading("6. Prioritized Recommendations", level=1)
    doc.add_heading("Tier 1 — Close the AOV Gap (Weeks 2–3)", level=2)
    add_table(doc,
        ["Action", "Target", "Est. Impact"],
        [
            ["POS hero ladder = online names/prices", "POS median $75", "+$40–69K/yr"],
            ["Delete price-named SKUs", "Enable ladder", "—"],
            ["Default card add-on at checkout", "25% attach", "+$3–5K/yr"],
        ],
    )

    doc.add_heading("Tier 2 — Close the Margin Gap (Weeks 4, 10)", level=2)
    add_table(doc,
        ["Action", "Target", "Est. Impact"],
        [
            ["Weekly stem-buy budget = forecast × 62%", "Wholesale under 45%", "+$8–15K GP/yr"],
            ["Raise Touch of Honey to $80–85", "38% hero GP", "Protect volume leader"],
            ["Recost top 8 SKUs", "Flag below 38%", "Pricing discipline"],
        ],
    )

    doc.add_heading("Tier 3 — Retention (Weeks 3, 7)", level=2)
    add_table(doc,
        ["Action", "Target", "Est. Impact"],
        [
            ["Merge 5,986 SS emails; occasion capture", "35% repeat in 90 days", "+$25–35K/yr at scale"],
            ["Sub offer to 54 past sub-buyers", "+10 active subs", "+$23K ARR"],
            ["Square email/phone on every tx", "70% capture in 90 days", "Unlocks CRM"],
        ],
    )

    doc.add_heading("7. 90-Day Targets", level=1)
    add_table(doc,
        ["Metric", "Baseline", "90-Day Target", "12-Month Target"],
        [
            ["POS median AOV", "$49", "$75", "$95+"],
            ["Blended AOV", "$78", "$90", "$110+"],
            ["Gross margin %", "31.5%", "38%", "45% → 62% path"],
            ["Wholesale % of revenue", "63%", "<55%", "<40%"],
            ["Card attach", "0.6%", "20%", "30%+"],
            ["Repeat rate (email)", "21%", "30%", "40%+"],
            ["Active subscriptions", "~0", "10", "+20% YoY"],
        ],
    )

    doc.add_heading("8. Monthly Tracking Checklist", level=1)
    add_bullets(doc, [
        "Revenue — Square net + SS total + bank deposits",
        "GP% — from accounts (never below 25% in any month)",
        "AOV — POS median, SS median, blended",
        "Online % — maintain 45–50%+",
        "Card attach % — target 25%+",
        "New vs returning email orders — target 45%+ returning",
        "Sub active count — manual until POS tracks",
    ])

    doc.add_heading("Appendix — Data Sources", level=1)
    add_table(doc,
        ["Data Point", "Source", "Confidence"],
        [
            ["Revenue, GP%, wholesale %", "Accounts forensic CSV", "High"],
            ["AOV, channel mix, seasonality", "Square + Squarespace exports", "High"],
            ["Product margins", "Weekly COGS allocation model", "Medium (predicted)"],
            ["Wedding / corporate", "Not in exports", "Low — needs tagging"],
            ["Benchmarks", "Industry benchmarks doc", "Reference only"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Dandelion Flowers · 1548 Webster St, Alameda, CA 94501 · (510) 522-2275 · "
        "Generated from commerce exports + accounts COGS · July 26, 2026"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build()
