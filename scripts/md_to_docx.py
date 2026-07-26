#!/usr/bin/env python3
"""Convert week-4-pl-analysis.md to a formatted Word document."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < ncols:
                cell = table.rows[i].cells[j]
                text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text.strip())
                text = text.replace("🔴", "").replace("🟡", "").replace("✓", "").replace("⚠️", "").strip()
                cell.text = text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                if i == 0:
                    set_cell_shading(cell, "E8F4E8")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
    doc.add_paragraph()


def add_rich_paragraph(doc, text, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)
    return p


def convert(md_path: Path, out_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = not in_code
            if not in_code:
                pass
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = "No Spacing"
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        if line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        if line.startswith("# "):
            t = doc.add_heading(line[2:].strip(), level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- [x] "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("\u2611  ")
            run.font.size = Pt(11)
            add_rich_paragraph_to(p, line[6:])
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("\u2610  ")
            run.font.size = Pt(11)
            add_rich_paragraph_to(p, line[6:])
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            add_rich_paragraph_to(p, text)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_paragraph_to(p, line[2:])
        elif line.strip().startswith("*") and line.strip().endswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.italic = True
            run.font.size = Pt(10)
        elif line.strip():
            add_rich_paragraph(doc, line.strip())
        i += 1

    if table_rows:
        add_table(doc, table_rows)

    doc.save(out_path)
    print(f"Saved {out_path}")


def add_rich_paragraph_to(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else root / "docs/week-4-pl-analysis.md"
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else root / "docs/Week_4_PL_Analysis.docx"
    convert(md, out)
